import requests
import docx
from docx.shared import Inches
from lxml import html
from selenium import webdriver
import commodity
import time
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains


def getData(url, outlet):
    ret=[]
    #build up internet connection
    options = webdriver.ChromeOptions()
    #options.add_argument('--ignore-certificate-errors')
    #options.add_argument('--headless')
    driver = webdriver.Chrome("C:/Users/Administrator/AppData/Local/Programs/Python/Python37/Scripts/chromedriver.exe", options=options)
    driver.get(url)
    
    #load more button
    if outlet:
        initialV=73
        buttonAddr='//*[@id="login-modal"]/div/div/button'
        try:
            WebDriverWait(driver, 90).until(EC.element_to_be_clickable((By.XPATH, '//*[@id="footer"]/div[6]/div[1]'))).click()
        except Exception as e:
            print(e)
       
    else:
        initialV=79
        buttonAddr='//*[@id="bx-element-1113062-R7KFV2S"]/button'
    
    #i=0
    initial=initialV
    time.sleep(10)
    while True:
        try:
            wait=WebDriverWait(driver, 120)
            if initial ==660:
                initial+=1
            wait.until(EC.presence_of_element_located((By.XPATH, f'//*[@id="cont"]/div[{initial}]/div')))
            element = driver.find_element(By.XPATH, f'//*[@id="cont"]/div[{initial}]/div')
            actions = ActionChains(driver)
            actions.move_to_element(element)
            actions.click(element)
            actions.perform()
        except Exception as e:
            try:
                WebDriverWait(driver, 30).until(EC.element_to_be_clickable((By.XPATH, buttonAddr))).click()
                initial=initialV
                continue
            except Exception as e2:
                print(e)
                print(e2)
                break
        #i+=1
        initial+=73
    
    page_source = driver.page_source
    driver.quit()
    tree=html.fromstring(page_source)
    
    #get all the links
    links=tree.xpath('//*[@id="cont"]/div/div/div/div/div[2]/h3/a/@href')
    print(f"number of links: {len(links)}")
    
    if outlet:
        pics=tree.xpath('//*[@id="cont"]/div/div/div/a[1]/img[1]/@src')
        print(f"number of pictures: {len(pics)}")
        prices=tree.xpath('//*[@id="cont"]/div/div/div/div/div[3]/div/div/div[2]/span[1]/span//text()')
        print(f"number of prices: {len(prices)}")
        names=tree.xpath('//*[@id="cont"]/div/div/div/div/div[2]/h3/a//text()')
        print(f"number of names: {len(names)}")
        
        minSize=min([len(links), len(pics), len(prices), len(names)])
        
    else:  
        pics=tree.xpath('//*[@id="cont"]/div/div/div[1]/a[1]/picture[1]/source[1]/@srcset')
        print(f"number of pictures: {len(pics)}")
        prices=[]
        for i in range(1400): #this number might need to change
        
            example1=tree.xpath(f'//*[@id="cont"]/div[{i}]/div/div[3]/div/div[3]/div/div/div/span[2]/span//text()')
            example2=tree.xpath(f'//*[@id="cont"]/div[{i}]/div/div[3]/div/div[3]/div/div/div//text()')
            example3=tree.xpath(f'//*[@id="cont"]/div[{i}]/div/div[3]/div/div[3]/div/div/span//text()')
            if len(example1)!=0:
                prices.append(example1[0])
            elif len(example2)!=0:
                prices.append(example2[0])
            elif len(example3)!=0:
                prices.append(example3[0])
        print(f"number of prices: {len(prices)}")
        names=tree.xpath('//*[@id="cont"]/div/div/div[3]/div/div[2]/h3/a//text()')
        print(f"number of names: {len(names)}")
        minSize=min([len(links), len(pics), len(prices), len(names)])
        
    
    
    
    for i in range(minSize):
        ret.append(commodity.commodity(names[i], pics[i], prices[i], links[i]))
    #for r in ret:
    #    print(r.name, r.pic, r.price, r.link)
    
    return ret #a list of commodity

def saveData(result, path):
    doc=docx.Document()
    for key in result:
        try:
            response = requests.get(key.pic.split()[0])
        except Exception as e:
            continue
        file = open("E:/sample_image.png", "wb")
        file.write(response.content)
        file.close()
        doc.add_picture("E:/sample_image.png", width=Inches(2.5))
        doc.add_paragraph(key.name+key.price)
        doc.add_paragraph(key.link)
        print("this is the coach: "+key.name+key.price+key.link)
        for content in result[key]:
            try:
                response = requests.get(content.pic)
            except Exception as e:
                continue
            file = open("E:/sample_image.png", "wb")
            file.write(response.content)
            file.close()
            doc.add_picture("E:/sample_image.png", width=Inches(2.5))
            doc.add_paragraph(content.name+ content.price)
            doc.add_paragraph(content.link + "similarity: " +str(content.similarity))
            print("this is the coach outlet: "+content.name+content.price+content.link)
        print("--------next one-------")
        doc.add_paragraph("--------------------------The end-------------------------------")
        doc.add_paragraph("--------------------------Another one---------------------------")
    doc.save(path)
    #os.remove("E:/sample_image.png")
    
    return  True

def main():
    #gather data
    #baseurl1=input("please enter a coach outlet page:")
    #baseurl2=input("please enter a coach page:")
    
    baseurl1="https://www.coachoutlet.com/shop/event-handbags-handbags-shoulder-bags?searchkeyword=Shoulder+Bags&qcat=text_header"
    baseurl2="https://www.coach.com/shop/women-handbags-shoulder-bags"
    
    coachOutlet=getData(baseurl1, True)
    coach=getData(baseurl2, False)
    print("data gathered")
       
    #analysis data
    result={}
    for i in coach:
        for j in coachOutlet:
            try:
                r = requests.post(
                    "https://api.deepai.org/api/image-similarity",
                    data={
                        'image1': i.pic.split()[0],
                        'image2': j.pic.split()[0]
                        },
                    headers={'api-key': '22411ab0-bf2a-4d67-9cd3-7b0d77fdc68b'})
            except Exception as e:
                print("delay")
                print(i.link)
                print(j.link)
                print()
                time.sleep(10)
                continue
            data=r.json()
            try:
                similarity=data['output']['distance']
                if similarity<15: # it's ok to compare the similarity and get the top 5 / do it only if the similarity is lower than x
                    j.similarity=similarity
                    if i in result:
                        result[i].append(j)
                    else:
                        result[i]=[j]
            except Exception as e:
                print(data)
                print(i.name+ " " +j.name + " skipped")
                print(i.pic.split()[0])
                print(j.pic.split()[0])
                continue
    
    print("analysis finished")
    
    
    #path=input("get the address where you want to store everything: ")
    path="E:/coachresult.docx"
    #save data
    if saveData(result, path):
        print("the result is ready")
    return


def help():
    print("hi")
    r = requests.post(
                "https://api.deepai.org/api/image-similarity",
                data={
                    'image1': "https://www.google.com/search?q=picture&source=lnms&tbm=isch&sa=X&ved=2ahUKEwjssMSFv-TpAhWPhHIEHTB9BJcQ_AUoAXoECCQQAw&biw=1280&bih=610&dpr=1.5",
                    'image2': "https://img2.cohimg.net/is/image/Coach/91022_imcah_a0?$plpMob$",
                },
                headers={'api-key': 'Your own API Key'})
    try:
        if r.json()['output']['distance']<=30:
            print("yes")
        else:
            print("no")
    except Exception as e:
        print(e)
    
    
    """
    //*[@id="cont"]/div[441]/div
    test=[]
    test.append(commodity.commodity("pig", "https://img2.cohimg.net/is/image/Coach/91022_imcah_a0?$plpMob$", "$45", "sss"))
    test.append(commodity.commodity("jhon", "https://img1.cohimg.net/is/image/Coach/29416_b4m2_a0?fmt=jpg&wid=680&hei=885&bgc=f0f0f0&fit=vfit&qlt=75", "46", "sssssss"))
    
    a=commodity.commodity("gssjjs", "https://img1.cohimg.net/is/image/Coach/651_v5quf_a0?fmt=jpg&wid=680&hei=885&bgc=f0f0f0&fit=vfit&qlt=75", "$7567", "sdfghj")
    b=commodity.commodity("sss", "https://img1.cohimg.net/is/image/Coach/79154_v5hgr_a0?fmt=jpg&wid=680&hei=885&bgc=f0f0f0&fit=vfit&qlt=75", "$23342", "sdwq")
    
    join={}
    join[a]=[b]
    join[b]=test
    
    path="E:/coachresult.docx"
    saveData(join, path)
    
    """
    """
    
    doc=docx.Document()
    doc.add_paragraph("hello")
    response = requests.get("https://img2.cohimg.net/is/image/Coach/91022_imcah_a0?$plpMob$")
    file = open("E:/sample_image.png", "wb")
    file.write(response.content)
    file.close()
    doc.add_picture("E:/sample_image.png")
    doc.save("E:/coachresult.docx")
    os.remove("E:/sample_image.png")
    
    """
    
    """
    options = webdriver.ChromeOptions()
    options.add_argument('--ignore-certificate-errors')
    #options.add_argument('--headless')
    driver = webdriver.Chrome("C:/Users/Administrator/AppData/Local/Programs/Python/Python37/Scripts/chromedriver.exe", options=options)
    driver.get("https://www.coach.com/coach-charlie-bucket-bag/55200.html?dwvar_color=V5QUF&cgid=women")
    
    #other colors
    try:
        WebDriverWait(driver, 90).until(EC.presence_of_element_located((By.CSS_SELECTOR, '#main > section.pdp-main.product-detail > div.product-information > div.product-info-master > div > div.title-area > div:nth-child(8) > ul')))
        color_list=driver.find_element(By.CSS_SELECTOR, '#main > section.pdp-main.product-detail > div.product-information > div.product-info-master > div > div.title-area > div:nth-child(8) > ul')
        
        colors = color_list.find_elements_by_tag_name("li")
    except Exception as e:
        try:
            WebDriverWait(driver, 30).until(EC.element_to_be_clickable((By.XPATH, '//*[@id="bx-element-1113062-R7KFV2S"]/button'))).click()
            WebDriverWait(driver, 90).until(EC.presence_of_element_located((By.CSS_SELECTOR, '#main > section.pdp-main.product-detail > div.product-information > div.product-info-master > div > div.title-area > div:nth-child(8) > ul')))
            color_list=driver.find_element(By.CSS_SELECTOR, '#main > section.pdp-main.product-detail > div.product-information > div.product-info-master > div > div.title-area > div:nth-child(8) > ul')
            colors = color_list.find_elements_by_tag_name("li")
        except Exception as e2:
            print(e)
            print(e2)
            driver.quit()
            return
        
    print(len(colors))  #0
    for color in colors:
        try:
            
            actions = ActionChains(driver)
            actions.move_to_element(color)
            actions.click(color)
            actions.perform()
            
            
        except Exception as e:
            try:
                WebDriverWait(driver, 30).until(EC.element_to_be_clickable((By.XPATH, '//*[@id="bx-element-1113062-R7KFV2S"]/button'))).click()
                continue
            except Exception as e2:
                print(e)
                print(e2)
                break
        page_source = driver.page_source
    
        tree=html.fromstring(page_source)
        pic=tree.xpath('//*[@id="pdp-carousel-element-1"]/img/@src')[0]
        price=tree.xpath('//*[@id="main"]/section[1]/div[2]/div[1]/div/div[1]/div[1]/div/span/span//text()')[-1].replace("\n", ""). replace(" ", ""). replace("\t", "")
        name=tree.xpath('//*[@id="main"]/section[1]/div[2]/div[1]/div/div[1]/h1//text()')[0]
    
    
        print(pic)
        print(price)
        print(name)
        #if [][][], stop
        
    
    
    driver.quit()
    """
main()
